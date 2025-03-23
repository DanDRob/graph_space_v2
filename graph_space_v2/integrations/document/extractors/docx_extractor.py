from typing import Dict, List, Any, BinaryIO
import os

from graph_space_v2.integrations.document.extractors.base import DocumentExtractor, DocumentInfo

# Use python-docx for DOCX processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxExtractor(DocumentExtractor):
    def extract(self, file_obj: BinaryIO, file_path: str) -> DocumentInfo:
        """Extract text and metadata from a DOCX document"""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX extraction but is not installed.")

        # Extract filename from path as fallback title
        filename = os.path.basename(file_path)
        base_filename = os.path.splitext(filename)[0]

        try:
            # Create a document object
            doc = docx.Document(file_obj)

            # Extract metadata
            metadata = {}
            core_properties = doc.core_properties
            if core_properties:
                if core_properties.title:
                    metadata['title'] = core_properties.title
                if core_properties.author:
                    metadata['author'] = core_properties.author
                if core_properties.created:
                    metadata['created'] = core_properties.created.isoformat()
                if core_properties.modified:
                    metadata['modified'] = core_properties.modified.isoformat()
                if core_properties.comments:
                    metadata['comments'] = core_properties.comments
                if core_properties.category:
                    metadata['category'] = core_properties.category
                if core_properties.subject:
                    metadata['subject'] = core_properties.subject
                if core_properties.keywords:
                    metadata['keywords'] = core_properties.keywords

            # Extract title from metadata or use filename
            title = metadata.get('title', base_filename)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text:
                    paragraphs.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))

            # Join all paragraphs with double newlines
            content = "\n\n".join(paragraphs)

            # Clean the extracted text
            content = self.clean_text(content)

            # If title is empty or None, try to extract it from content
            if not title:
                title = self.extract_title_from_text(content, base_filename)

            # Count pages (approximation since docx doesn't have direct page count)
            # Estimate based on typical page having ~3000 characters
            estimated_pages = max(1, len(content) // 3000)

            return DocumentInfo(
                title=title,
                content=content,
                metadata=metadata,
                file_path=file_path,
                file_type="docx",
                pages=estimated_pages
            )

        except Exception as e:
            # Handle extraction errors
            return DocumentInfo(
                title=base_filename,
                content=f"Error extracting DOCX content: {str(e)}",
                metadata={"error": str(e)},
                file_path=file_path,
                file_type="docx",
                pages=0
            )
